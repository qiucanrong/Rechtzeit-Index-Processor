from __future__ import annotations

import io
from typing import List, Tuple

import streamlit as st
from docx import Document

from index_processor import process_lines, process_text

def read_txt(uploaded_bytes: bytes) -> str:
    try:
        return uploaded_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return uploaded_bytes.decode("cp1252")

def read_docx(uploaded_bytes: bytes) -> List[str]:
    """
    Returns a list of strings, one per paragraph, WITHOUT trailing '\n'.
    We'll treat each paragraph as a "line".
    """
    bio = io.BytesIO(uploaded_bytes)
    doc = Document(bio)
    return [p.text for p in doc.paragraphs]


def write_docx(lines: List[str]) -> bytes:
    """
    lines: list of output lines, each likely ends with '\n'
    Writes each line as a paragraph (without trailing newline).
    """
    doc = Document()
    for line in lines:
        doc.add_paragraph(line.rstrip("\n"))
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


st.set_page_config(page_title="Rechtzeit Index Processor", layout="centered")
st.title("Judaica Division Rechtzeit Index Processor")

uploaded = st.file_uploader("**Upload Raw Index File (.docx or .txt)**", type=["txt", "docx"])

output_format = st.radio(
    "**Output Format**",
    options=["docx", "txt"],
    horizontal=True,
)

if uploaded is not None:
    file_name = uploaded.name
    raw_bytes = uploaded.getvalue()

    # Read input as list of lines (strings without '\n' for docx), or as raw text for txt
    if file_name.lower().endswith(".txt"):
        raw_text = read_txt(raw_bytes)
        processed_text = process_text(raw_text)
        processed_lines = processed_text.splitlines(True)
    else:
        # docx
        input_paras = read_docx(raw_bytes)
        # convert to lines w/ '\n' so processor behavior is consistent
        processed_lines = process_lines([p + "\n" for p in input_paras])
        processed_text = "".join(processed_lines)

    if output_format == "txt":
        st.download_button(
            label="**Download Processed File** (.txt)",
            data=processed_text.encode("utf-8"),
            file_name="index_processed.txt",
            mime="text/plain",
        )
    else:
        docx_bytes = write_docx(processed_lines)
        st.download_button(
            label="**Download Processed File** (.docx)",
            data=docx_bytes,
            file_name="index_processed.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )